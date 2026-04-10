# -*- coding: utf-8 -*-
"""
生成越南语情景对话 Word 文件
页边距: 1.2cm (上下左右)
页脚: Apex Personal Workspace © 2026 (居中，每页底部)
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import sys

def set_run_font(run, font_name='Arial', font_name_east='微软雅黑', size=12):
    """设置文本字体（支持中文）"""
    run.font.size = Pt(size)
    run.font.name = font_name
    # 设置东亚字体（中文）
    r = run._element.rPr
    if r is None:
        r = OxmlElement('w:rPr')
        run._element.insert(0, r)
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name_east)

def add_page_footer(doc, text):
    """添加页脚到每一页"""
    section = doc.sections[0]
    
    # 获取页脚并确保可见
    footer = section.footer
    section.footer.is_linked_to_previous = False
    
    # 清除现有内容
    for p in footer.paragraphs:
        p.clear()
    
    # 设置居中文本
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run(text)
    set_run_font(run, 'Arial', '微软雅黑', 9)
    
    # 确保页脚可见
    sectPr = section._sectPr
    titlePg = sectPr.find(qn('w:titlePg'))
    if titlePg is not None:
        titlePg.getparent().remove(titlePg)

def create_word(dialogues, output_path):
    """创建 Word 文档"""
    doc = Document()
    
    # 设置页边距为 1.2cm
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    
    # 添加页脚
    add_page_footer(doc, 'Apex Personal Workspace © 2026')
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('越南语情景对话')
    set_run_font(run, 'Arial', '微软雅黑', 18)
    run.font.bold = True
    
    # 导出日期
    from datetime import datetime
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f'导出日期: {datetime.now().strftime("%Y-%m-%d")}')
    set_run_font(run, 'Arial', '微软雅黑', 10)
    
    # 统一分隔线
    separator = '─' * 60
    
    # 分隔线
    sep = doc.add_paragraph()
    run = sep.add_run(separator)
    set_run_font(run, 'Arial', '微软雅黑', 10)
    
    # 添加对话内容
    for d in dialogues:
        # 日期
        date_p = doc.add_paragraph()
        run = date_p.add_run(f'日期: {d.get("date", "")}')
        set_run_font(run, 'Arial', '微软雅黑', 10)
        run.font.bold = True
        
        # 主题
        topic_p = doc.add_paragraph()
        run = topic_p.add_run(f'主题: {d.get("topic", "")}')
        set_run_font(run, 'Arial', '微软雅黑', 11)
        run.font.bold = True
        
        # 内容 - 中文和越南语放在同一行
        content = d.get('content', '')
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line:
                # 检查下一行是否存在
                next_line = lines[i+1].strip() if i+1 < len(lines) else ''
                
                # 判断当前行是越南语还是中文
                is_vn = any(c in line for c in 'ăâđêôơưàằầèềìòồờùừỳáắấéếíóốớúứýảẳẩẻểỉỏổởủửỷạặậẹệịọộợụựỵ')
                next_is_vn = any(c in next_line for c in 'ăâđêôơưàằầèềìòồờùừỳáắấéếíóốớúứýảẳẩẻểỉỏổởủửỷạặậẹệịọộợụựỵ') if next_line else False
                
                # 如果当前是越南语且下一行是中文，合并为一行
                if is_vn and next_line and not next_is_vn:
                    p = doc.add_paragraph()
                    run = p.add_run(f'{line}  |  {next_line}')
                    set_run_font(run, 'Arial', '微软雅黑', 12)
                    i += 2
                    continue
                # 如果当前是中文且下一行是越南语，合并为一行
                elif not is_vn and next_line and next_is_vn:
                    p = doc.add_paragraph()
                    run = p.add_run(f'{next_line}  |  {line}')
                    set_run_font(run, 'Arial', '微软雅黑', 12)
                    i += 2
                    continue
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    set_run_font(run, 'Arial', '微软雅黑', 12)
            i += 1
        
        # 分隔
        doc.add_paragraph()
        sep = doc.add_paragraph()
        run = sep.add_run(separator)
        set_run_font(run, 'Arial', '微软雅黑', 10)
    
    # 保存
    doc.save(output_path)
    print(f'Word 文件已保存: {output_path}')

if __name__ == '__main__':
    # 示例数据
    dialogues = [
        {
            'date': '2026-04-10',
            'topic': 'goi do uong o quan ca phe - 咖啡店点单',
            'content': '''A: Xin chao! Ban muon dung gi a?
你好！你想要用什么？

B: Cho toi mot ca phe sua da it duong.
给我一杯冰牛奶咖啡少糖。

A: Vang, xin quy khach cho mot chut.
好的，请稍等。

B: Cam on ban.
谢谢你。'''
        },
        {
            'date': '2026-04-10',
            'topic': 'o san bay chuan bi len may bay - 机场登机',
            'content': '''A: Xin chao, toi muon lam thu tuc len may bay.
你好，我想办理登机手续。

B: Cho toi xem ho chieu cua anh.
请让我看一下您的护照。

A: Day a.
这是。

B: Anh ngoi ghe 12A, cua len may bay so 5.
您坐在12A座位，登机口5号。'''
        }
    ]
    
    output = 'vietnamese_dialogues.docx'
    if len(sys.argv) > 1:
        output = sys.argv[1]
    
    create_word(dialogues, output)
